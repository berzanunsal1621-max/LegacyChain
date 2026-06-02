# -*- coding: utf-8 -*-
"""
LegacyChain - Teknofest 2026 Blokzincir Yarışması
Ön Değerlendirme Raporu (ÖDR) Word Dosyası Oluşturucu
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# AYARLAR
# ============================================================

OUTPUT_FILE = "LegacyChain_ODR_2026.docx"

def create_document():
    doc = Document()
    
    # -- Sayfa kenar boşlukları: 2.5 cm --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # -- Varsayılan stil ayarları --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Arial fontunu East Asian ve Complex Script için de ayarla
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:cs'), 'Arial')
    
    return doc


def add_heading_black(doc, text, level=1):
    """Arial Black, 14pt, koyu mavi başlık ekler"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial Black'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 70, 127)  # Koyu mavi (şablondaki renk)
    run.bold = True
    # Fonts
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Arial Black')
    rFonts.set(qn('w:hAnsi'), 'Arial Black')
    rFonts.set(qn('w:cs'), 'Arial Black')
    return p


def add_sub_heading(doc, text):
    """Alt başlık - Arial Black, 12pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Arial Black'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 70, 127)
    run.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Arial Black')
    rFonts.set(qn('w:hAnsi'), 'Arial Black')
    rFonts.set(qn('w:cs'), 'Arial Black')
    return p


def add_body(doc, text):
    """Normal gövde metni - Arial, 12pt, iki yana yaslı"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Madde imi ile metin"""
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.name = 'Arial'
        run_bold.font.size = Pt(12)
        run_bold.bold = True
        run_normal = p.add_run(text)
        run_normal.font.name = 'Arial'
        run_normal.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    return p


def add_table(doc, headers, rows):
    """Formatlanmış tablo ekler"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Başlık satırı
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = True
        # Arka plan rengi
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '00467F')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Veri satırları
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            # Alternatif satır renklendirme
            if row_idx % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8F0FE')
                shading.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading)
    
    doc.add_paragraph()  # Tablo sonrası boşluk
    return table


def add_page_break(doc):
    doc.add_page_break()


# ============================================================
# RAPOR İÇERİĞİ
# ============================================================

def build_cover_page(doc):
    """Kapak Sayfası"""
    # Boşluk
    for _ in range(4):
        doc.add_paragraph()
    
    # Başlık
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BLOKZİNCİR YARIŞMASI")
    run.font.name = 'Arial Black'
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 70, 127)
    run.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("ÖN DEĞERLENDİRME RAPORU")
    run2.font.name = 'Arial Black'
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0, 70, 127)
    run2.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Proje bilgileri
    fields = [
        ("Proje Adı:", "LegacyChain – Merkeziyetsiz Dijital Varlık Miras Protokolü"),
        ("Takım Adı:", "[TAKIMADI]"),
        ("Takım ID:", "[TAKIMID]"),
        ("Başvuru ID:", "[BAŞVURUID]"),
        ("Seçilen Yarışma Kapsamı:", "Blokzincir Yarışması"),
    ]
    
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        run_label = p.add_run(label + " ")
        run_label.font.name = 'Arial Black'
        run_label.font.size = Pt(14)
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.name = 'Arial'
        run_value.font.size = Pt(12)
    
    add_page_break(doc)


def build_toc(doc):
    """İçindekiler Sayfası"""
    add_heading_black(doc, "İÇİNDEKİLER")
    doc.add_paragraph()
    
    items = [
        "1.  PROJE ÖZETİ",
        "2.  KATMA DEĞER VE YENİLİKÇİLİK",
        "3.  TEKNOLOJİ KULLANIMI",
        "4.  UYGULANABİLİRLİK",
        "5.  YAYGIN ETKİ",
        "6.  SÜRDÜRÜLEBİLİRLİK",
        "7.  PROJE TAKVİMİ",
        "8.  TAKIM YAPISI",
        "9.  KAYNAKÇA",
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    
    add_page_break(doc)


def build_section1(doc):
    """1. PROJE ÖZETİ"""
    add_heading_black(doc, "1. PROJE ÖZETİ")
    
    add_body(doc,
        "Kripto para ve dijital varlık ekosisteminin hızla büyümesiyle birlikte, varlık sahiplerinin vefatı "
        "veya uzun süreli erişimsizliği durumunda dijital varlıkların yasal mirasçılara güvenli biçimde "
        "aktarılması kritik bir problem haline gelmiştir. Günümüzde tahminen 140 milyar dolar değerinde "
        "Bitcoin, sahiplerinin özel anahtarlarını miras bırakamaması nedeniyle kalıcı olarak erişilemez "
        "durumdadır [1]. Geleneksel miras sistemleri ise avukat, noter ve mahkeme gibi merkezi aracılara "
        "bağımlı olup yüksek maliyet, bürokratik gecikme ve sınır ötesi yetki karmaşası gibi sorunlarla "
        "karşı karşıyadır [2]."
    )
    
    add_body(doc,
        "LegacyChain, Ethereum blokzinciri üzerinde çalışan merkeziyetsiz bir dijital varlık miras "
        "protokolüdür. Projenin temel amacı, kripto varlık sahiplerinin varlıklarını birden fazla mirasçıya "
        "yüzdelik paylarla, herhangi bir merkezi otoriteye gerek duymaksızın güvenli ve otomatik biçimde "
        "aktarabilmesini sağlamaktır."
    )
    
    add_body(doc, "Protokol, dört temel mekanizma üzerine inşa edilmiştir:")
    
    add_bullet(doc,
        " Varlık sahibi, belirli aralıklarla zincir üstü ping() işlemi göndererek aktif olduğunu "
        "kanıtlar. Ayrıca Ghost Ping (Otonom Aktivite Takibi) mekanizması, sahip herhangi bir zincir "
        "üstü işlem yaptığında bu aktiviteyi otomatik olarak kaydeder; böylece manuel ping ihtiyacını azaltır.",
        "Proof-of-Life (Yaşam Kanıtı) Protokolü:"
    )
    
    add_bullet(doc,
        " Birden fazla bağımsız otorite (hastane, resmi kurum, aile üyesi gibi), sahibin vefat durumunu "
        "zincir üstünde imzalar. 2/3 konsensüs sağlandığında miras süreci tetiklenir.",
        "Merkeziyetsiz Oracle Konsensüsü:"
    )
    
    add_bullet(doc,
        " Miras tetiklendikten sonra 24 saatlik bir bekleme süresi uygulanır. Bu süre zarfında eğer sahip "
        "hayatta ise işlemi iptal edebilir; böylece yanlış alarm senaryolarına karşı koruma sağlanır.",
        "Grace Period (Güvenlik Süresi):"
    )
    
    add_bullet(doc,
        " Sistemde en fazla 10 mirasçı tanımlanabilir ve her birine yüzdelik pay atanır. Süre dolduğunda "
        "akıllı kontrat, fonları tanımlanan oranlarda otomatik olarak dağıtır.",
        "Çoklu Varis Desteği ve Otomatik Dağıtım:"
    )
    
    add_body(doc,
        "Proje, Solidity ^0.8.0 ile geliştirilen akıllı kontratlar, React tabanlı bir DApp arayüzü ve "
        "Hardhat test altyapısı kullanılarak hayata geçirilmiştir. Ethereum Sepolia test ağı üzerinde "
        "başarıyla deploy edilmiş ve 40 otomatik test senaryosuyla doğrulanmıştır."
    )


def build_section2(doc):
    """2. KATMA DEĞER VE YENİLİKÇİLİK"""
    add_heading_black(doc, "2. KATMA DEĞER VE YENİLİKÇİLİK")
    
    add_sub_heading(doc, "2.1 Problem Tanımı")
    
    add_body(doc,
        "Kripto para piyasasının toplam değeri 2025 itibarıyla 2 trilyon doları aşmış olup, bireysel "
        "yatırımcıların dijital varlıklarını güvenle miras bırakmaları konusunda ciddi bir altyapı "
        "eksikliği bulunmaktadır [3]. Geleneksel miras sistemi ile dijital varlık dünyası arasındaki "
        "temel uyumsuzluklar aşağıdaki tabloda özetlenmiştir:"
    )

    add_table(doc,
        ["Problem", "Geleneksel Çözüm", "Kısıtlama"],
        [
            ["Erişim kontrolü", "Avukat, banka", "Merkezi, yavaş, pahalı"],
            ["Doğrulama", "Ölüm sertifikası", "Manuel, bürokratik gecikmeler"],
            ["Sınır ötesi varlık", "Uluslararası mahkemeler", "Yetki alanı sorunları"],
            ["Gizlilik", "Resmi kayıtlar", "Varlık bilgilerinin ifşası"],
        ]
    )
    
    add_sub_heading(doc, "2.2 Mevcut Blokzincir Çözümlerinin Eksikleri")
    
    add_body(doc,
        "Halihazırda piyasada bulunan blokzincir tabanlı miras çözümleri incelendiğinde önemli "
        "eksiklikler göze çarpmaktadır [4]:"
    )
    
    add_bullet(doc,
        " Pek çok mevcut çözüm yalnızca tek bir mirasçıya dağıtım yapabilmektedir. Gerçek dünya "
        "senaryolarında ise bir bireyin varlıklarını eş, çocuk ve diğer aile bireyleri arasında "
        "paylaştırması en yaygın ihtiyaçtır.",
        "Tekil mirasçı sınırlaması:"
    )
    
    add_bullet(doc,
        " Varis değişikliklerinde bekleme süresi uygulanmadığından, bir saldırgan cüzdan sahibini "
        "baskı altında anlık değişiklik yapmaya zorlayabilir.",
        "Zaman kilidi eksikliği:"
    )
    
    add_bullet(doc,
        " Tek bir oracle kaynağına bağımlılık, tek nokta arızası (single point of failure) riski oluşturur.",
        "Merkezi oracle bağımlılığı:"
    )
    
    add_bullet(doc,
        " Varlıklar ya tamamıyla ya da hiç dağıtılmakta; yüzdelik paylaştırma mümkün olmamaktadır.",
        "Kısmi miras desteği yokluğu:"
    )
    
    add_sub_heading(doc, "2.3 LegacyChain'in Farkı")
    
    add_body(doc,
        "LegacyChain, yukarıdaki eksiklikleri aşan kapsamlı bir çözüm sunmaktadır. Aşağıdaki tabloda "
        "mevcut çözümlerle karşılaştırmalı analiz sunulmuştur:"
    )
    
    add_table(doc,
        ["Özellik", "LegacyChain", "Mevcut Çözüm A", "Mevcut Çözüm B"],
        [
            ["Çoklu varis desteği", "✅ 10'a kadar", "❌ Tekil", "✅ 3'e kadar"],
            ["Yüzdelik dağıtım", "✅ Özelleştirilebilir", "❌ Yok", "⚠️ Sabit"],
            ["Zaman kilidi (TimeLock)", "✅ 24 saat", "❌ Yok", "✅ 48 saat"],
            ["ReentrancyGuard", "✅", "❌", "✅"],
            ["Multi-sig acil durum", "✅ 2/3", "❌", "❌"],
            ["Oracle entegrasyonu", "✅ Çoklu", "✅ Tekil", "❌"],
            ["Grace Period", "✅ 24 saat", "❌", "❌"],
            ["Ghost Ping", "✅", "❌", "❌"],
            ["Açık kaynak", "✅", "❌", "✅"],
        ]
    )
    
    add_body(doc, "Özgün katkılar:")
    
    add_bullet(doc,
        " Sahip akıllı kontrat ile herhangi bir etkileşimde bulunduğunda aktivite otomatik olarak "
        "kaydedilir. Bu, kullanıcının sürekli olarak ping() çağrısı yapmasını zorunlu kılmaz ve "
        "kullanıcı deneyimini önemli ölçüde iyileştirir. Bu yaklaşım, incelenen mevcut çözümlerin "
        "hiçbirinde bulunmamaktadır.",
        "Ghost Ping Mekanizması:"
    )
    
    add_bullet(doc,
        " Tek bir oracle'a güvenmek yerine, birden fazla bağımsız otoriteden 2/3 çoğunluk konsensüsü "
        "gerektiren bir doğrulama sistemi geliştirilmiştir. Bu, hem güvenilirliği artırır hem de merkezi "
        "bir başarısızlık noktası riskini ortadan kaldırır.",
        "Merkeziyetsiz Oracle Konsensüsü:"
    )
    
    add_bullet(doc,
        " ReentrancyGuard, TimeLock ve Multi-sig mekanizmalarının bir arada kullanılması, protokole "
        "çok katmanlı bir güvenlik yapısı kazandırır. Bu bütünleşik yaklaşım, incelenen blokzincir "
        "miras projelerinde nadir görülmektedir.",
        "Katmanlı Güvenlik Modeli:"
    )
    
    add_bullet(doc,
        " Proje, Türkiye'deki blokzincir ekosisteminde dijital miras alanında özgün bir çözüm olarak "
        "konumlanmaktadır. Türkiye'de kripto varlık kullanımının hızla artmasıyla birlikte [5], bu alanda "
        "yerli bir çözümün geliştirilmesi hem teknolojik bağımsızlık hem de yerel mevzuata uyum açısından "
        "stratejik bir önem taşımaktadır.",
        "Yerlilik ve Özgünlük:"
    )


def build_section3(doc):
    """3. TEKNOLOJİ KULLANIMI"""
    add_heading_black(doc, "3. TEKNOLOJİ KULLANIMI")
    
    add_sub_heading(doc, "3.1 Sistem Mimarisi")
    
    add_body(doc, "LegacyChain, üç katmanlı bir mimariye sahiptir:")
    
    add_bullet(doc,
        " Tüm iş mantığı Ethereum blokzinciri üzerinde çalışan akıllı kontratlarda işlenmektedir. Bu "
        "katman, miras yönetimi, zamanlayıcı kontrolü, oracle sinyalleri ve fon dağıtımını kapsamaktadır.",
        "Katman 1 – Akıllı Kontrat Katmanı (On-Chain):"
    )
    
    add_bullet(doc,
        " Merkeziyetsiz oracle kontratı, birden fazla yetkili otoriteden gelen vefat sinyallerini toplar "
        "ve 2/3 konsensüs sağlandığında ana kontrata bildirim yapar.",
        "Katman 2 – Oracle Katmanı:"
    )
    
    add_bullet(doc,
        " React tabanlı DApp, MetaMask cüzdan entegrasyonu ile kullanıcıların akıllı kontrat ile "
        "etkileşim kurmasını sağlar.",
        "Katman 3 – Kullanıcı Arayüzü Katmanı (Off-Chain):"
    )
    
    add_sub_heading(doc, "3.2 Akıllı Kontrat Tasarımı")
    
    add_body(doc, "Proje kapsamında dört adet Solidity akıllı kontratı geliştirilmiştir:")
    
    add_table(doc,
        ["Kontrat", "Satır Sayısı", "İşlev"],
        [
            ["MultiHeirInheritance.sol", "577", "Ana vault kontratı – çoklu varis yönetimi, TimeLock, oracle entegrasyonu, ReentrancyGuard, Grace Period"],
            ["DecentralizedOracle.sol", "94", "Oracle kontratı – çoklu otorite yönetimi, konsensüs tabanlı vefat doğrulama (2/3 imza)"],
            ["Inheritance.sol", "71", "Basit kullanım senaryoları için hafif, tek varisli versiyon"],
            ["MockERC20.sol", "~50", "ERC-20 token miras senaryolarının test edilmesi için sahte token kontratı"],
        ]
    )
    
    add_sub_heading(doc, "3.2.1 Temel Veri Yapıları")
    
    add_body(doc,
        "Varis bilgileri, Heir struct'ı ile temsil edilmektedir. Her mirasçı için cüzdan adresi, "
        "miras yüzdesi (1-100 arası), tanımlayıcı isim, aktiflik durumu ve Commit-Reveal mekanizması "
        "için secretHash alanları tutulmaktadır. TimeLock mekanizması için bekleyen değişiklikler ise "
        "PendingChange struct'ında yeni mirasçı adresi, yeni yüzde, yeni isim, kilit açılma zamanı "
        "ve var olma durumu bilgileriyle saklanmaktadır."
    )
    
    add_sub_heading(doc, "3.2.2 Temel Fonksiyonlar")
    
    add_table(doc,
        ["Fonksiyon", "Erişim", "Açıklama"],
        [
            ["ping()", "Sahip", "Proof-of-Life zamanlayıcısını sıfırlar"],
            ["recordActivity()", "Sahip", "Ghost Ping – otonom aktivite kaydı"],
            ["addHeir()", "Sahip", "Yeni mirasçı ekler (otomatik yüzde dengeleme)"],
            ["initiateHeirUpdate()", "Sahip", "TimeLock'lu güncelleme başlatır"],
            ["executeHeirUpdate()", "Sahip", "24 saat sonra güncellemeyi uygular"],
            ["simulateOracleSignal()", "Oracle", "Vefat sinyali gönderir"],
            ["startGracePeriod()", "Herkes", "Grace Period süresini başlatır"],
            ["claimInheritance()", "Herkes", "Fonları yüzdelik olarak dağıtır"],
            ["claimTokens()", "Herkes", "ERC-20 tokenları dağıtır"],
            ["emergencyWithdraw()", "Sahip", "Acil geri çekim"],
            ["emergencyMultiSigTransfer()", "2/3 İmza", "Multi-sig acil transfer"],
        ]
    )
    
    add_sub_heading(doc, "3.3 Güvenlik Mekanizmaları")
    
    add_body(doc,
        "3.3.1 ReentrancyGuard: Yeniden giriş saldırılarına karşı koruma sağlamak amacıyla OpenZeppelin'in "
        "ReentrancyGuard kalıbı uygulanmıştır [6]. Bu mekanizma, fon transferi gerçekleştiren tüm fonksiyonlarda "
        "(claimInheritance, emergencyWithdraw, emergencyMultiSigTransfer, claimTokens) aktiftir. Özellikle "
        "2016'daki The DAO saldırısında [7] kullanılan reentrancy saldırı vektörüne karşı doğrudan koruma sağlamaktadır."
    )
    
    add_body(doc,
        "3.3.2 TimeLock Koruma: Varis değişikliklerinde 24 saatlik zorunlu bekleme süresi uygulanır. Bu mekanizma "
        "üç kritik senaryoya karşı koruma sağlar: baskı (coercion) saldırıları, hata kurtarma ve mirasçı "
        "bildirimi. Saldırgan, sahib anlık değişiklik yapmaya zorlasa bile 24 saatlik süre koruma penceresi "
        "oluşturur. Yanlışlıkla yapılan değişiklikler uygulama öncesinde iptal edilebilir."
    )
    
    add_body(doc,
        "3.3.3 Multi-Sig Acil Durum Protokolü: Kritik acil durum operasyonları için 2/3 çoklu imza mekanizması "
        "uygulanmıştır. Üç güvenilir imzacı tanımlanmıştır: kontrat sahibi, oracle adresi ve birincil mirasçı. "
        "Bu yapı, herhangi bir tek tarafın fonları kötüye kullanmasını engeller."
    )
    
    add_body(doc,
        "3.3.4 Erişim Kontrolü: Rol tabanlı erişim kontrol mekanizması onlyOwner, onlyOracle ve "
        "onlyTrustedSigner modifier'ları ile uygulanmıştır. Her fonksiyon, yalnızca yetkili roller tarafından "
        "çağrılabilmektedir."
    )
    
    add_sub_heading(doc, "3.4 Teknoloji Yığını")
    
    add_table(doc,
        ["Bileşen", "Teknoloji"],
        [
            ["Blokzincir", "Ethereum (Sepolia Test Ağı)"],
            ["Akıllı Kontrat Dili", "Solidity ^0.8.0"],
            ["Ön Yüz", "React (CDN) + Tailwind CSS + Ethers.js v5.7"],
            ["Test Altyapısı", "Hardhat + Chai"],
            ["Cüzdan Entegrasyonu", "MetaMask"],
            ["Güvenlik", "ReentrancyGuard, TimeLock, Multi-sig"],
        ]
    )
    
    add_sub_heading(doc, "3.5 Test ve Doğrulama")
    
    add_body(doc,
        "Proje kapsamında 40 otomatik test senaryosu geliştirilmiş ve tamamı başarıyla geçmiştir:"
    )
    
    add_table(doc,
        ["Kategori", "Test Sayısı", "Durum"],
        [
            ["Erişim kontrolü", "4", "✅ Başarılı"],
            ["Çoklu varis yönetimi", "5", "✅ Başarılı"],
            ["TimeLock mekanizması", "4", "✅ Başarılı"],
            ["Miras dağıtım", "4", "✅ Başarılı"],
            ["Reentrancy koruması", "2", "✅ Başarılı"],
            ["Zamanlayıcı mantığı", "4", "✅ Başarılı"],
            ["Sınır durumları", "4", "✅ Başarılı"],
            ["Gas optimizasyonu", "3", "✅ Başarılı"],
            ["Gelişmiş protokol testleri", "10", "✅ Başarılı"],
            ["Toplam", "40", "✅ Tamamı Başarılı"],
        ]
    )
    
    add_sub_heading(doc, "3.5.1 Gas Optimizasyon Sonuçları")
    
    add_body(doc,
        "Akıllı kontrat işlemlerinin gas maliyetleri ölçülmüş ve hedef değerlerin altında kalmıştır:"
    )
    
    add_table(doc,
        ["İşlem", "Harcanan Gas", "Hedef", "Durum"],
        [
            ["ping()", "29.728", "< 50.000", "✅"],
            ["addHeir()", "155.797", "< 200.000", "✅"],
            ["claimInheritance() (3 varis)", "91.262", "< 300.000", "✅"],
        ]
    )
    
    add_body(doc,
        "Düşük gas tüketimi, protokolün gerçek dünya kullanımında ekonomik olarak sürdürülebilir olmasını "
        "sağlamaktadır."
    )


def build_section4(doc):
    """4. UYGULANABİLİRLİK"""
    add_heading_black(doc, "4. UYGULANABİLİRLİK")
    
    add_sub_heading(doc, "4.1 Projenin Hayata Geçirilme Planı")
    
    add_body(doc,
        "LegacyChain'in prototip aşamasından ticari ürüne dönüştürülmesi aşağıdaki yol haritası ile "
        "planlanmaktadır:"
    )
    
    add_bullet(doc,
        " Çoklu varis akıllı kontratı, merkeziyetsiz oracle kontratı, kapsamlı test süiti ve temel DApp "
        "arayüzü tamamlanmıştır. Sistem, Ethereum Sepolia test ağında çalışır durumdadır.",
        "Aşama 1 – Prototip (Mevcut Durum) ✅:"
    )
    
    add_bullet(doc,
        " ERC-20 token ve NFT (ERC-721) miras desteğinin eklenmesiyle protokolün kapsadığı varlık türleri "
        "genişletilecektir. Bu, DeFi tokenleri, stablecoin'ler ve dijital koleksiyonlar gibi farklı varlık "
        "sınıflarını miras sürecine dahil edecektir.",
        "Aşama 2 – Genişletilmiş Varlık Desteği:"
    )
    
    add_bullet(doc,
        " Chainlink gibi endüstri standardı merkeziyetsiz oracle ağlarıyla entegrasyon sağlanacaktır [8]. "
        "Bu sayede vefat doğrulama süreci, gerçek dünya veri kaynaklarıyla desteklenecektir.",
        "Aşama 3 – Gerçek Oracle Entegrasyonu:"
    )
    
    add_bullet(doc,
        " Ethereum dışında Polygon, BNB Smart Chain ve diğer EVM uyumlu zincirlere genişleme ile "
        "erişilebilirlik artırılacaktır.",
        "Aşama 4 – Çoklu Zincir Desteği:"
    )
    
    add_sub_heading(doc, "4.2 Ticarileşme Potansiyeli")
    
    add_body(doc,
        "LegacyChain'in ticarileşme modeli üç gelir kanalı üzerine yapılandırılmıştır:"
    )
    
    add_bullet(doc,
        " Kullanıcıların akıllı kontrat deploy etmesi sırasında tek seferlik bir hizmet bedeli.",
        "Vault oluşturma ücreti:"
    )
    
    add_bullet(doc,
        " Gelişmiş oracle entegrasyonu, çoklu zincir desteği ve kurumsal düzeyde raporlama gibi ek "
        "özellikler abonelik modeliyle sunulabilir.",
        "Premium özellikler:"
    )
    
    add_bullet(doc,
        " Kripto borsaları, dijital varlık saklama hizmeti sunan kuruluşlar ve sigorta şirketleri için "
        "özelleştirilmiş B2B miras çözümleri.",
        "Kurumsal çözümler:"
    )
    
    add_sub_heading(doc, "4.3 Hedef Kullanıcı Senaryoları")
    
    add_body(doc,
        "Senaryo 1 – Aile İçi Miras: Bir kripto varlık sahibi, varlıklarını eşine (%50), birinci çocuğuna "
        "(%25) ve ikinci çocuğuna (%25) paylaştırmak ister. LegacyChain ile bu dağıtım otomatik ve güvenli "
        "biçimde gerçekleşir."
    )
    
    add_body(doc,
        "Senaryo 2 – İş Sürekliliği: Bir startup kurucusu, şirketin kripto hazinesini ortak kurucuya (%40), "
        "CFO'ya (%30) ve hukuk danışmanına (%30) paylaştırarak, kendi erişimsizliği durumunda iş "
        "operasyonlarının kesintisiz sürmesini sağlar."
    )
    
    add_body(doc,
        "Senaryo 3 – Hayırseverlik: Bir birey, varlıklarının %50'sini hayır kurumuna, %50'sini ailesine "
        "bırakarak, hayırseverlik bağışının ölümden sonra da garanti altına alınmasını sağlar."
    )


def build_section5(doc):
    """5. YAYGIN ETKİ"""
    add_heading_black(doc, "5. YAYGIN ETKİ")
    
    add_sub_heading(doc, "5.1 Toplumsal Etki")
    
    add_body(doc,
        "Kripto para kullanıcı sayısı dünya genelinde 500 milyonu aşmış olup [9], Türkiye kripto para "
        "sahipliğinde Avrupa'da lider konumdadır [5]. Bu büyüme, dijital miras sorununu her geçen gün "
        "daha geniş kitleleri etkileyen bir mesele haline getirmektedir. LegacyChain şu toplumsal etkileri "
        "hedeflemektedir:"
    )
    
    add_bullet(doc,
        " Geleneksel miras sistemine erişimi olmayan bireyler (bankasız nüfus) için düşük maliyetli, "
        "merkeziyetsiz bir alternatif sunar.",
        "Finansal kapsayıcılık:"
    )
    
    add_bullet(doc,
        " Dijital varlık sahiplerinin ailelerinin ekonomik güvencesini zincir üstü garanti altına alır.",
        "Aile güvencesi:"
    )
    
    add_bullet(doc,
        " Blokzincir teknolojisinin somut ve anlaşılır bir kullanım alanı sunarak toplumun bu "
        "teknolojiyle tanışmasını kolaylaştırır.",
        "Dijital okur-yazarlık:"
    )
    
    add_sub_heading(doc, "5.2 Ekonomik Etki")
    
    add_bullet(doc,
        " 140 milyar dolarlık kayıp Bitcoin probleminin çözümüne doğrudan katkı sağlar [1].",
        "Kayıp varlık sorununun azaltılması:"
    )
    
    add_bullet(doc,
        " Avukat, noter ve banka gibi aracılara ödenen yüksek komisyonların yerine düşük gas maliyetli "
        "zincir üstü işlemler sunar.",
        "Aracı maliyetlerinin ortadan kaldırılması:"
    )
    
    add_bullet(doc,
        " Türkiye'nin blokzincir ekosisteminde özgün bir DeFi ürünü olarak yerli teknoloji kapasitesine "
        "katkı sağlar.",
        "Yerli teknoloji geliştirme:"
    )
    
    add_sub_heading(doc, "5.3 Endüstri Etkisi")
    
    add_bullet(doc, " Borsa müşterilerine entegre miras hizmeti sunma imkânı.", "Kripto borsaları:")
    add_bullet(doc, " Saklama hizmetlerine otomatik miras aktarım katmanı ekleme.", "Dijital varlık saklama kuruluşları:")
    add_bullet(doc, " Dijital varlık sigortası ürünleriyle entegrasyon potansiyeli.", "Sigorta sektörü:")


def build_section6(doc):
    """6. SÜRDÜRÜLEBİLİRLİK"""
    add_heading_black(doc, "6. SÜRDÜRÜLEBİLİRLİK")
    
    add_sub_heading(doc, "6.1 Finansal Sürdürülebilirlik")
    
    add_body(doc,
        "LegacyChain, uzun vadede finansal açıdan sürdürülebilir bir model üzerine kurgulanmıştır:"
    )
    
    add_bullet(doc,
        " Akıllı kontratlar bir kez deploy edildikten sonra blokzincir üzerinde otonom olarak çalışır. "
        "Sunucu veya bakım maliyeti gerektirmez.",
        "Düşük operasyonel maliyet:"
    )
    
    add_bullet(doc,
        " Protokolün tüm temel işlemleri gas hedeflerinin altında gerçekleştirilmekte olup, kullanıcılar "
        "için ekonomik erişim sağlanmaktadır.",
        "Gas optimizasyonu:"
    )
    
    add_bullet(doc,
        " Vault oluşturma ücretleri ve premium özellik abonelikleri ile sürdürülebilir bir gelir akışı oluşturulabilir.",
        "Gelir modeli:"
    )
    
    add_sub_heading(doc, "6.2 Teknik Sürdürülebilirlik")
    
    add_bullet(doc,
        " Proje tamamen açık kaynak olup, topluluk katkısına açıktır. Bu yaklaşım, bağımsız güvenlik "
        "denetimleri ve sürekli iyileştirme olanağı sağlar.",
        "Açık kaynak:"
    )
    
    add_bullet(doc,
        " Akıllı kontratlar modüler yapıda tasarlanmıştır. Yeni özellikler (ERC-20, NFT, çapraz zincir) "
        "mevcut sistemi bozmadan eklenebilir.",
        "Modüler mimari:"
    )
    
    add_bullet(doc,
        " Ethereum'un sürekli gelişen altyapısı üzerine inşa edilmiş olması, uzun vadeli teknolojik uyumu "
        "garanti eder [10].",
        "Ethereum ekosistemi:"
    )
    
    add_sub_heading(doc, "6.3 Çevresel Sürdürülebilirlik")
    
    add_body(doc,
        "Ethereum'un Proof-of-Stake (PoS) konsensüs mekanizmasına geçişiyle (The Merge, Eylül 2022), "
        "ağın enerji tüketimi %99,95 oranında azalmıştır [11]. LegacyChain, PoS tabanlı Ethereum üzerinde "
        "çalıştığından çevresel ayak izi minimum düzeydedir."
    )
    
    add_sub_heading(doc, "6.4 Risk Yönetimi")
    
    add_table(doc,
        ["Risk", "Etki", "Azaltma Stratejisi"],
        [
            ["Akıllı kontrat açığı", "Yüksek", "Kapsamlı test süiti (40 test), ReentrancyGuard, TimeLock"],
            ["Yanlış Oracle sinyali", "Yüksek", "2/3 konsensüs mekanizması, Grace Period (24 saat)"],
            ["Ethereum ağ sorunları", "Orta", "Çoklu zincir desteği planı (Polygon, BSC)"],
            ["Regülasyon değişiklikleri", "Orta", "Modüler mimari ile uyum güncellemesi kapasitesi"],
            ["Özel anahtar kaybı", "Yüksek", "Multi-sig acil durum protokolü (2/3 onay)"],
        ]
    )


def build_section7(doc):
    """7. PROJE TAKVİMİ"""
    add_heading_black(doc, "7. PROJE TAKVİMİ")
    
    add_body(doc,
        "Proje kapsamında iş paketleri alt faaliyetleri ile birlikte aşağıdaki tabloda sunulmuştur:"
    )
    
    add_table(doc,
        ["İş Paketi / Alt Faaliyet", "Başlangıç", "Bitiş", "Durum"],
        [
            ["İP1: Araştırma ve Tasarım", "", "", ""],
            ["  Literatür taraması ve mevcut çözüm analizi", "Ocak 2026", "Ocak 2026", "✅ Tamamlandı"],
            ["  Sistem mimarisi tasarımı", "Ocak 2026", "Şubat 2026", "✅ Tamamlandı"],
            ["  Protokol spesifikasyonu (Whitepaper)", "Şubat 2026", "Şubat 2026", "✅ Tamamlandı"],
            ["İP2: Akıllı Kontrat Geliştirme", "", "", ""],
            ["  Temel miras kontratı (Inheritance.sol)", "Şubat 2026", "Şubat 2026", "✅ Tamamlandı"],
            ["  Çoklu varis kontratı (MultiHeirInheritance.sol)", "Şubat 2026", "Mart 2026", "✅ Tamamlandı"],
            ["  Oracle kontratı (DecentralizedOracle.sol)", "Mart 2026", "Mart 2026", "✅ Tamamlandı"],
            ["İP3: Güvenlik ve Test", "", "", ""],
            ["  Birim testleri (40 test senaryosu)", "Mart 2026", "Mart 2026", "✅ Tamamlandı"],
            ["  Gas optimizasyonu ve performans testleri", "Mart 2026", "Mart 2026", "✅ Tamamlandı"],
            ["  Güvenlik açığı taraması", "Mart 2026", "Nisan 2026", "🔄 Devam ediyor"],
            ["İP4: Kullanıcı Arayüzü", "", "", ""],
            ["  React DApp arayüzü geliştirme", "Mart 2026", "Mart 2026", "✅ Tamamlandı"],
            ["  MetaMask cüzdan entegrasyonu", "Mart 2026", "Mart 2026", "✅ Tamamlandı"],
            ["  Kullanıcı deneyimi iyileştirmeleri", "Nisan 2026", "Mayıs 2026", "📋 Planlandı"],
            ["İP5: İleri Özellikler", "", "", ""],
            ["  ERC-20 token miras desteği", "Mayıs 2026", "Haziran 2026", "📋 Planlandı"],
            ["  NFT miras desteği", "Haziran 2026", "Temmuz 2026", "📋 Planlandı"],
            ["  Chainlink oracle entegrasyonu", "Temmuz 2026", "Ağustos 2026", "📋 Planlandı"],
            ["İP6: Yayınlama ve Belgelendirme", "", "", ""],
            ["  Sepolia test ağı deploy", "Mart 2026", "Mart 2026", "✅ Tamamlandı"],
            ["  Kapsamlı dokümantasyon", "Nisan 2026", "Mayıs 2026", "🔄 Devam ediyor"],
            ["  Mainnet hazırlık ve denetim", "Ağustos 2026", "Eylül 2026", "📋 Planlandı"],
        ]
    )


def build_section8(doc):
    """8. TAKIM YAPISI"""
    add_heading_black(doc, "8. TAKIM YAPISI")
    
    add_body(doc,
        "[Bu bölümü kendi takım yapınıza göre doldurunuz. Aşağıda örnek bir şablon verilmiştir.]"
    )
    
    add_table(doc,
        ["Rol", "Uzmanlık Alanı", "Sorumluluklar"],
        [
            ["Takım Lideri / Akıllı Kontrat Geliştiricisi", "Solidity, Ethereum, Akıllı Kontrat Güvenliği", "Akıllı kontrat tasarımı ve geliştirme, güvenlik mekanizmalarının implementasyonu, test senaryolarının yazılması"],
            ["[Üye 2 Rolü]", "[Uzmanlık]", "[Sorumluluklar]"],
            ["[Üye 3 Rolü]", "[Uzmanlık]", "[Sorumluluklar]"],
        ]
    )
    
    add_body(doc,
        "Not: Takım üyelerinin isim ve fotoğraf gibi kişisel bilgileri raporda yer almamaktadır. "
        "(Şablon kuralları gereği)"
    )


def build_section9(doc):
    """9. KAYNAKÇA"""
    add_heading_black(doc, "9. KAYNAKÇA")
    
    references = [
        '[1] Chainalysis, "The 2024 Crypto Crime Report: Lost and Stolen Cryptocurrency," 2024, Erişim Tarihi: Mart 2026, https://www.chainalysis.com/blog/crypto-crime-report/',
        '[2] Zetzsche, D. A., Buckley, R. P., Arner, D. W. (2020) "Decentralized Finance," Journal of Financial Regulation, Cilt 6, Sayı 2, s. 172–203, DOI: 10.1093/jfr/fjaa010.',
        '[3] CoinMarketCap, "Total Cryptocurrency Market Capitalization," 2025, Erişim Tarihi: Mart 2026, https://coinmarketcap.com/charts/',
        '[4] Buterin, V. (2014) "Ethereum: A Next-Generation Smart Contract and Decentralized Application Platform," Ethereum Whitepaper, https://ethereum.org/en/whitepaper/',
        '[5] Triple-A, "Global Crypto Ownership Data 2024," 2024, Erişim Tarihi: Mart 2026, https://triple-a.io/crypto-ownership-data/',
        '[6] OpenZeppelin, "ReentrancyGuard Documentation," Erişim Tarihi: Mart 2026, https://docs.openzeppelin.com/contracts/4.x/api/security#ReentrancyGuard',
        '[7] Mehar, M. I., Shier, C. L., Giambattista, A. (2019) "Understanding a Revolutionary and Flawed Grand Experiment in Blockchain: The DAO Attack," Journal of Cases on Information Technology, Cilt 21, Sayı 1, s. 19–32, DOI: 10.4018/JCIT.2019010102.',
        '[8] Chainlink, "Chainlink Oracle Network Architecture," Erişim Tarihi: Mart 2026, https://docs.chain.link/architecture-overview/architecture-overview',
        '[9] Crypto.com, "Crypto Market Sizing Report," Ocak 2025, Erişim Tarihi: Mart 2026, https://crypto.com/research',
        '[10] Ethereum Foundation, "Ethereum Improvement Proposals (EIPs)," Erişim Tarihi: Mart 2026, https://eips.ethereum.org/',
        '[11] Ethereum Foundation, "The Merge," Eylül 2022, Erişim Tarihi: Mart 2026, https://ethereum.org/en/roadmap/merge/',
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        run.font.name = 'Arial'
        run.font.size = Pt(11)


# ============================================================
# ANA FONKSİYON
# ============================================================

def main():
    print("LegacyChain ÖDR Word dosyası oluşturuluyor...")
    
    doc = create_document()
    
    # 1. Kapak Sayfası
    build_cover_page(doc)
    
    # 2. İçindekiler
    build_toc(doc)
    
    # 3. Bölümler
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)
    build_section5(doc)
    build_section6(doc)
    build_section7(doc)
    build_section8(doc)
    
    # Kaynakça ayrı sayfada
    add_page_break(doc)
    build_section9(doc)
    
    # Kaydet
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, OUTPUT_FILE)
    doc.save(output_path)
    print(f"✅ Rapor başarıyla oluşturuldu: {output_path}")
    print()
    print("📝 Doldurmanız gereken yerler:")
    print("   - Kapak: [TAKIMADI], [TAKIMID], [BAŞVURUID]")
    print("   - Bölüm 8: Takım üyelerinin rolleri ve sorumlulukları")


if __name__ == "__main__":
    main()
